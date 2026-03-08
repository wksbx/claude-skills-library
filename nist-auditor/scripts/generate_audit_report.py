"""
Generate an Audit Summary Report in Word format from audit findings.

Usage:
    python generate_audit_report.py --findings findings.json --output audit_report.docx

Input: JSON file with findings and metadata
Output: Formatted Word document audit report
"""

import json
import sys
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Error: python-docx required. Install with: pip install python-docx --break-system-packages")
    sys.exit(1)


RISK_COLORS = {
    "Critical": RGBColor(255, 0, 0),
    "High": RGBColor(255, 102, 0),
    "Medium": RGBColor(255, 215, 0),
    "Low": RGBColor(0, 176, 80),
}


def get_risk_level(score):
    if score >= 16: return "Critical"
    elif score >= 10: return "High"
    elif score >= 5: return "Medium"
    return "Low"


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def generate_report(data, output_path):
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading(level=0)
    run = title.add_run("NIST Security Assessment Report")
    run.font.size = Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = data.get("metadata", {})
    entity = meta.get("entity", "Organization")
    framework = meta.get("framework", "NIST SP 800-53 Rev 5")
    assessor = meta.get("assessor", "TBD")
    date = meta.get("date", datetime.now().strftime("%Y-%m-%d"))
    baseline = meta.get("baseline", "Moderate")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Entity: {entity}\n").bold = True
    p.add_run(f"Framework: {framework}\n")
    p.add_run(f"Baseline: {baseline}\n")
    p.add_run(f"Assessment Date: {date}\n")
    p.add_run(f"Assessor: {assessor}")

    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Scope and Methodology",
        "3. Findings Summary",
        "4. Detailed Findings",
        "5. Risk Distribution",
        "6. Recommendations",
        "7. Next Steps",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)

    findings = data.get("findings", [])
    total = len(findings)
    critical = sum(1 for f in findings if get_risk_level(f.get("risk_score", 0)) == "Critical")
    high = sum(1 for f in findings if get_risk_level(f.get("risk_score", 0)) == "High")
    medium = sum(1 for f in findings if get_risk_level(f.get("risk_score", 0)) == "Medium")
    low = sum(1 for f in findings if get_risk_level(f.get("risk_score", 0)) == "Low")

    exec_summary = meta.get("executive_summary", "")
    if exec_summary:
        doc.add_paragraph(exec_summary)
    else:
        doc.add_paragraph(
            f"This report presents the findings from a {framework} security assessment "
            f"conducted for {entity}. The assessment evaluated the implementation and "
            f"effectiveness of security controls against the {baseline} baseline."
        )

    doc.add_paragraph(
        f"The assessment identified {total} findings: "
        f"{critical} Critical, {high} High, {medium} Medium, and {low} Low risk."
    )

    # Summary table
    table = doc.add_table(rows=5, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [("Risk Level", "Count")]
    rows_data = [
        ("Critical", str(critical)),
        ("High", str(high)),
        ("Medium", str(medium)),
        ("Low", str(low)),
    ]

    table.cell(0, 0).text = "Risk Level"
    table.cell(0, 1).text = "Count"
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

    risk_hex = {"Critical": "FF0000", "High": "FF6600", "Medium": "FFD700", "Low": "00B050"}
    for i, (level, count) in enumerate(rows_data, 1):
        table.cell(i, 0).text = level
        table.cell(i, 1).text = count
        set_cell_shading(table.cell(i, 0), risk_hex[level])

    # 2. Scope and Methodology
    doc.add_heading("2. Scope and Methodology", level=1)

    scope_text = meta.get("scope", "")
    if scope_text:
        doc.add_paragraph(scope_text)
    else:
        doc.add_paragraph(
            f"The assessment scope included all {baseline} baseline controls for the "
            f"{entity} information system as defined by {framework}."
        )

    doc.add_heading("Assessment Methods", level=2)
    doc.add_paragraph("The assessment employed three methods per SP 800-53A:")
    methods = [
        "Examine — Review of policies, procedures, plans, configurations, and logs",
        "Interview — Discussions with system owners, administrators, and security personnel",
        "Test — Technical testing of security controls including vulnerability scanning and configuration verification",
    ]
    for m in methods:
        doc.add_paragraph(m, style="List Bullet")

    # 3. Findings Summary
    doc.add_heading("3. Findings Summary", level=1)

    # By family
    families = {}
    for f in findings:
        family = f.get("control_family", "Unknown")
        if family not in families:
            families[family] = []
        families[family].append(f)

    if families:
        table = doc.add_table(rows=1 + len(families), cols=4)
        table.style = "Light Grid Accent 1"
        for i, header in enumerate(["Control Family", "Findings", "Highest Risk", "Key Issues"]):
            table.cell(0, i).text = header
            for p in table.cell(0, i).paragraphs:
                for run in p.runs:
                    run.font.bold = True

        for i, (family, family_findings) in enumerate(sorted(families.items()), 1):
            highest = max(f.get("risk_score", 0) for f in family_findings)
            table.cell(i, 0).text = family
            table.cell(i, 1).text = str(len(family_findings))
            table.cell(i, 2).text = get_risk_level(highest)
            issues = "; ".join(f.get("title", "") for f in family_findings[:3])
            table.cell(i, 3).text = issues

    # 4. Detailed Findings
    doc.add_heading("4. Detailed Findings", level=1)

    sorted_findings = sorted(findings, key=lambda f: f.get("risk_score", 0), reverse=True)
    for idx, finding in enumerate(sorted_findings, 1):
        risk_score = finding.get("risk_score", 0)
        risk_level = get_risk_level(risk_score)

        doc.add_heading(f"Finding {idx}: {finding.get('title', 'Untitled')}", level=2)

        # Finding details table
        details = [
            ("Control ID(s)", finding.get("control_ids", "N/A")),
            ("Control Family", finding.get("control_family", "N/A")),
            ("Status", finding.get("status", "N/A")),
            ("Risk Score", f"{risk_score} ({risk_level})"),
            ("Likelihood", str(finding.get("likelihood", "N/A"))),
            ("Impact", str(finding.get("impact", "N/A"))),
        ]

        table = doc.add_table(rows=len(details), cols=2)
        table.style = "Light Grid Accent 1"
        for i, (label, value) in enumerate(details):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            for p in table.cell(i, 0).paragraphs:
                for run in p.runs:
                    run.font.bold = True

        if finding.get("description"):
            doc.add_heading("Description", level=3)
            doc.add_paragraph(finding["description"])

        if finding.get("evidence"):
            doc.add_heading("Evidence", level=3)
            doc.add_paragraph(finding["evidence"])

        if finding.get("remediation"):
            doc.add_heading("Recommended Remediation", level=3)
            doc.add_paragraph(finding["remediation"])

    # 5. Risk Distribution
    doc.add_heading("5. Risk Distribution", level=1)
    doc.add_paragraph(
        f"The following summarizes the risk distribution across all {total} findings:"
    )

    distribution = [
        f"Critical (16-25): {critical} findings — Immediate action required",
        f"High (10-15): {high} findings — Remediate within 90 days",
        f"Medium (5-9): {medium} findings — Remediate within 180 days",
        f"Low (1-4): {low} findings — Monitor and address as resources permit",
    ]
    for d in distribution:
        doc.add_paragraph(d, style="List Bullet")

    # 6. Recommendations
    doc.add_heading("6. Recommendations", level=1)

    recommendations = data.get("recommendations", [])
    if recommendations:
        for rec in recommendations:
            doc.add_paragraph(rec, style="List Number")
    else:
        doc.add_paragraph(
            "Recommendations will be provided based on the detailed findings above. "
            "Priority should be given to Critical and High risk findings."
        )

    # 7. Next Steps
    doc.add_heading("7. Next Steps", level=1)
    next_steps = data.get("next_steps", [
        "Review and validate findings with system owners",
        "Develop Plan of Action & Milestones (POA&M) for all open findings",
        "Assign remediation owners and establish timelines",
        "Schedule follow-up assessment to verify remediation effectiveness",
        "Update System Security Plan (SSP) to reflect current control status",
    ])
    for step in next_steps:
        doc.add_paragraph(step, style="List Number")

    doc.save(output_path)
    print(f"Audit report saved to: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Generate NIST Audit Summary Report")
    parser.add_argument("--findings", required=True, help="Path to findings JSON file")
    parser.add_argument("--output", default="audit_report.docx", help="Output Word file path")
    args = parser.parse_args()

    with open(args.findings, "r") as f:
        data = json.load(f)

    generate_report(data, args.output)


if __name__ == "__main__":
    main()
